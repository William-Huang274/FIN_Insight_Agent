import asyncio
import io
from uuid import uuid4

import pytest

from sec_agent.research_foundation.task_attachments import TaskAttachmentStore, parse_document, MAX_BYTES
from sec_agent.research_foundation.source_document_navigation import SourceDocumentRequest


def request(store, thread, **kwargs):
    return asyncio.run(store.read(thread_id=thread, request=SourceDocumentRequest(source_space="uploads", **kwargs)))


def test_upload_markdown_structure_search_read_and_task_isolation(tmp_path):
    store = TaskAttachmentStore(tmp_path)
    thread = str(uuid4())
    added = store.add(thread, "supply.md", b"# Supply\n\n## Memory\nDRAM cost rose 12 percent.\n\n## Capacity\nProduction is still ramping.")
    catalog = request(store, thread, operation="catalog")
    assert catalog.total_matches == 1 and not catalog.items[0]["writer_citable"]
    result = request(store, thread, operation="search", query="DRAM", document_id=added["document_id"])
    assert len(result.items) == 1 and "DRAM" in result.items[0]["preview"]
    read = request(store, thread, operation="read", document_id=added["document_id"], node_id=result.items[0]["node_id"])
    assert read.items[0]["writer_citable"] and not read.items[0]["numeric_fact_authority"]
    assert read.items[0]["source_role"] == "user_upload_unverified"
    assert str(tmp_path) not in read.model_dump_json()
    with pytest.raises(ValueError, match="not_in_current_task"):
        request(store, str(uuid4()), operation="read", document_id=added["document_id"])
    assert TaskAttachmentStore(tmp_path).list(thread)[0]["document_id"] == added["document_id"]


@pytest.mark.parametrize("name,data", [("../outside.txt", b"x"), ("C:\\private.txt", b"x"), ("x.svg", b"svg"),
    ("script.exe", b"MZ"), ("x.pdf", b"not a PDF"), ("x.txt", b"a\x00b"), ("x.txt", b"")])
def test_invalid_uploads_are_rejected_without_rows(tmp_path, name, data):
    store, thread = TaskAttachmentStore(tmp_path), str(uuid4())
    with pytest.raises(ValueError):
        store.add(thread, name, data)
    assert not store.list(thread)


def test_html_scripts_not_executed_or_indexed_and_table_order_retained():
    pages, _ = parse_document("financials.html", b"<h1>Results</h1><script>stealSecret()</script><table><tr><th>Period</th><th>Revenue</th></tr><tr><td>2025</td><td>100</td></tr></table><p>Unaudited</p>")
    text = "\n".join(p["text"] for p in pages)
    assert "stealSecret" not in text and "Period | Revenue" in text and text.index("100") < text.index("Unaudited")


def test_real_pdf_pages_and_docx_tables(tmp_path):
    from reportlab.pdfgen import canvas
    from docx import Document
    stream = io.BytesIO()
    pdf = canvas.Canvas(stream)
    pdf.drawString(60, 730, "Revenue 100 million USD")
    pdf.showPage()
    pdf.drawString(60, 730, "Operating profit 12 million USD")
    pdf.save()
    store, thread = TaskAttachmentStore(tmp_path), str(uuid4())
    item = store.add(thread, "earnings.pdf", stream.getvalue())
    page = request(store, thread, operation="read", document_id=item["document_id"], page_start=2)
    assert "Operating profit" in page.items[0]["passage"] and page.items[0]["parser_page_start"] == 2
    assert store.image(thread, item["document_id"], 2).startswith(b"\x89PNG")
    doc = Document()
    doc.add_heading("Quarterly results", 1)
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Revenue"
    table.cell(0, 1).text = "100"
    doc.add_paragraph("Footnote after table")
    stream = io.BytesIO()
    doc.save(stream)
    pages, _ = parse_document("results.docx", stream.getvalue())
    assert "Revenue | 100" in pages[0]["text"]
    assert pages[0]["text"].index("100") < pages[0]["text"].index("Footnote")


def test_image_inspection_source_binding_and_native_cache(tmp_path):
    from PIL import Image
    stream = io.BytesIO()
    Image.new("RGB", (320, 200), "white").save(stream, format="PNG")
    store, thread = TaskAttachmentStore(tmp_path), str(uuid4())
    item = store.add(thread, "chart.png", stream.getvalue())
    assert item["needs_vision"]
    with pytest.raises(ValueError, match="inspect_image"):
        request(store, thread, operation="read", document_id=item["document_id"])
    calls = []
    async def model(image, question):
        calls.append(question)
        return "Test image has no numbers; cannot infer any revenue."
    req = SourceDocumentRequest(source_space="uploads", operation="inspect_image", document_id=item["document_id"])
    result = asyncio.run(store.read(thread_id=thread, request=req, vision_reader=model))
    again = asyncio.run(store.read(thread_id=thread, request=req, vision_reader=model))
    assert result == again and len(calls) == 1
    assert "user_upload_vision" in result.items[0]["source_role"]
    assert not result.numeric_fact_authority
